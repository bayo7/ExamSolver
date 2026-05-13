"""
parser/docx_parser.py
---------------------
.docx dosyasını okuyup her MCQ sorusunu yapısal bir dict'e dönüştürür.
Inline ve anchored görselleri base64 olarak yakalar.
"""

import re
import base64
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from PIL import Image
import io
from typing import List, Dict, Any


# Soru başlangıcını tanıyan regex kalıpları
# Örnekler: "1.", "Q1.", "Question 1:", "Soru 1 -"
QUESTION_PATTERNS = [
    r"^(Q\s*\d+[\.\:\-\)])",       # Q1. Q1: Q1) Q 1.
    r"^(Question\s+\d+[\.\:\-\)])", # Question 1. Question 1:
    r"^(Soru\s*\d+[\.\:\-\)])",     # Soru 1. Soru 1:
    r"^(\d+[\.\)])\s+\S",           # 1. veya 1) + boşluk + karakter
]

# Şık başlangıcını tanıyan regex kalıpları
OPTION_PATTERNS = [
    r"^([A-D])[\.\)\-\:]\s+",       # A. B) C- D:
    r"^\(([A-D])\)\s+",             # (A) (B)
]


def is_question_start(text: str) -> bool:
    """Metnin yeni bir soru başlangıcı olup olmadığını kontrol eder."""
    text = text.strip()
    for pattern in QUESTION_PATTERNS:
        if re.match(pattern, text, re.IGNORECASE):
            return True
    return False


def parse_option(text: str) -> tuple[str | None, str]:
    """
    Şık satırından harf ve içeriği ayırır.
    Örnek: "A. Paris" -> ("A", "Paris")
    Şık değilse -> (None, text)
    """
    text = text.strip()
    for pattern in OPTION_PATTERNS:
        m = re.match(pattern, text)
        if m:
            letter = m.group(1).upper()
            content = text[m.end():].strip()
            return letter, content
    return None, text


def extract_images_from_element(element, doc) -> List[Dict]:
    """
    Bir XML elementinden (paragraph veya run) tüm görselleri çeker.
    Hem inline hem anchored drawing'leri yakalar.
    """
    images = []

    # XML namespace'leri
    NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    NS_A  = "http://schemas.openxmlformats.org/drawingml/2006/main"
    NS_R  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    # Hem inline hem anchor içindeki blip'leri bul
    for blip in element.iter("{%s}blip" % NS_A):
        rId = blip.get("{%s}embed" % NS_R)
        if rId and rId in doc.part.related_parts:
            image_part = doc.part.related_parts[rId]
            raw = image_part.blob

            # Pillow ile format doğrula, gerekirse PNG'e çevir
            try:
                img = Image.open(io.BytesIO(raw))
                img_format = img.format or "PNG"
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                buf.seek(0)
                b64 = base64.b64encode(buf.read()).decode("utf-8")
                images.append({
                    "data": b64,
                    "format": "png",
                    "media_type": "image/png",
                    "width": img.width,
                    "height": img.height,
                })
            except Exception as e:
                # Pillow açamazsa ham base64 gönder
                b64 = base64.b64encode(raw).decode("utf-8")
                mime = image_part.content_type
                images.append({
                    "data": b64,
                    "format": mime.split("/")[-1],
                    "media_type": mime,
                    "width": None,
                    "height": None,
                })

    return images


def parse_exam(docx_path: str) -> List[Dict[str, Any]]:
    """
    Ana parser fonksiyonu.
    .docx dosyasını okur ve soru listesi döner.

    Dönen yapı:
    [
      {
        "question_number": 1,
        "question_text": "...",
        "options": {"A": "...", "B": "...", "C": "...", "D": "..."},
        "images": [{"data": "<base64>", "format": "png", ...}]
      },
      ...
    ]
    """
    doc = Document(docx_path)
    questions: List[Dict[str, Any]] = []

    current_q: Dict[str, Any] | None = None
    q_number = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        # Bu paragraftaki görselleri çek
        para_images = extract_images_from_element(para._element, doc)

        if not text and not para_images:
            continue  # Boş satır, atla

        # --- Yeni soru mu başlıyor? ---
        if is_question_start(text):
            # Önceki soruyu kaydet
            if current_q is not None:
                questions.append(current_q)

            q_number += 1
            # Soru numarasını ve önekini temizle
            clean_text = re.sub(
                r"^(Q\s*\d+|Question\s+\d+|Soru\s*\d+|\d+)[\.\:\-\)]\s*",
                "", text, flags=re.IGNORECASE
            ).strip()

            current_q = {
                "question_number": q_number,
                "question_text": clean_text,
                "options": {},
                "images": para_images,  # Soruyla birlikte gelen görseller
            }

        # --- Şık satırı mı? ---
        elif current_q is not None:
            letter, content = parse_option(text)
            if letter:
                current_q["options"][letter] = content
                # Şıkla birlikte gelen görseli şık'a ekle
                if para_images:
                    current_q["options"][letter + "_images"] = para_images
                    current_q["images"].extend(para_images)
            else:
                # Ne soru başlangıcı ne şık → soru metnine ekle (çok satırlı sorular)
                if text:
                    current_q["question_text"] += " " + text
                if para_images:
                    current_q["images"].extend(para_images)

    # Son soruyu da ekle
    if current_q is not None:
        questions.append(current_q)

    # Tablolardaki soruları da tara (bazı sınavlar tablo kullanır)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if is_question_start(cell_text):
                    # Tablo hücresinden de ayrıştır (basit versiyon)
                    q_number += 1
                    clean_text = re.sub(
                        r"^(Q\s*\d+|Question\s+\d+|Soru\s*\d+|\d+)[\.\:\-\)]\s*",
                        "", cell_text, flags=re.IGNORECASE
                    ).strip()
                    cell_images = extract_images_from_element(cell._element, doc)
                    questions.append({
                        "question_number": q_number,
                        "question_text": clean_text,
                        "options": {},
                        "images": cell_images,
                    })

    print(f"[Parser] {len(questions)} soru bulundu → {docx_path}")
    return questions


if __name__ == "__main__":
    # Hızlı test
    import sys, json
    if len(sys.argv) < 2:
        print("Kullanım: python docx_parser.py <dosya.docx>")
        sys.exit(1)
    result = parse_exam(sys.argv[1])
    for q in result:
        print(f"\nQ{q['question_number']}: {q['question_text'][:80]}...")
        for k, v in q['options'].items():
            if not k.endswith("_images"):
                print(f"  {k}) {v}")
        print(f"  Görsel sayısı: {len(q['images'])}")
