"""
sample_papers/create_sample_exams.py
-------------------------------------
Test için iki örnek MCQ sınav dosyası üretir:
  1. math_exam.docx   — Matematik soruları
  2. science_exam.docx — Fen/Bilim soruları

Her dosya >= 10 soru içerir.
Cevap anahtarı da ayrıca kaydedilir (doğruluk testi için).
"""

import json
from docx import Document
from docx.shared import Pt
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent


# ─────────────────────────────────────────────
# Soru verileri
# ─────────────────────────────────────────────

MATH_QUESTIONS = [
    ("What is the value of √144?",
     {"A": "10", "B": "11", "C": "12", "D": "14"}, "C"),

    ("If 3x + 7 = 22, what is x?",
     {"A": "3", "B": "4", "C": "5", "D": "6"}, "C"),

    ("What is 15% of 200?",
     {"A": "25", "B": "30", "C": "35", "D": "40"}, "B"),

    ("A triangle has angles 45°, 90°, and x°. Find x.",
     {"A": "35", "B": "40", "C": "45", "D": "55"}, "C"),

    ("What is the area of a circle with radius 7? (Use π ≈ 3.14)",
     {"A": "153.86", "B": "43.96", "C": "78.5", "D": "154"}, "A"),

    ("Simplify: 2³ × 2²",
     {"A": "2⁴", "B": "2⁵", "C": "2⁶", "D": "4⁵"}, "B"),

    ("What is the slope of the line y = 3x − 5?",
     {"A": "-5", "B": "5", "C": "3", "D": "-3"}, "C"),

    ("Find the LCM of 12 and 18.",
     {"A": "6", "B": "24", "C": "36", "D": "72"}, "C"),

    ("If a = 4 and b = 3, what is a² + b²?",
     {"A": "25", "B": "7", "C": "49", "D": "14"}, "A"),

    ("What is the median of: 3, 7, 5, 9, 1?",
     {"A": "3", "B": "5", "C": "7", "D": "9"}, "B"),

    ("How many degrees are in a full rotation?",
     {"A": "180", "B": "270", "C": "360", "D": "90"}, "C"),

    ("What is 0.25 as a fraction in simplest form?",
     {"A": "1/2", "B": "1/4", "C": "2/5", "D": "3/8"}, "B"),
]


SCIENCE_QUESTIONS = [
    ("Which gas do plants absorb during photosynthesis?",
     {"A": "Oxygen", "B": "Nitrogen", "C": "Carbon Dioxide", "D": "Hydrogen"}, "C"),

    ("What is the chemical formula for water?",
     {"A": "CO₂", "B": "H₂O₂", "C": "HO", "D": "H₂O"}, "D"),

    ("Which planet is closest to the Sun?",
     {"A": "Venus", "B": "Mars", "C": "Mercury", "D": "Earth"}, "C"),

    ("What is the powerhouse of the cell?",
     {"A": "Nucleus", "B": "Ribosome", "C": "Mitochondria", "D": "Golgi Apparatus"}, "C"),

    ("What force pulls objects toward Earth?",
     {"A": "Magnetism", "B": "Friction", "C": "Tension", "D": "Gravity"}, "D"),

    ("Which element has the atomic number 1?",
     {"A": "Helium", "B": "Hydrogen", "C": "Lithium", "D": "Carbon"}, "B"),

    ("What type of rock is formed from cooled lava?",
     {"A": "Sedimentary", "B": "Metamorphic", "C": "Igneous", "D": "Fossil"}, "C"),

    ("The speed of light is approximately:",
     {"A": "300 km/s", "B": "3,000 km/s", "C": "300,000 km/s", "D": "30,000 km/s"}, "C"),

    ("Which vitamin is produced when skin is exposed to sunlight?",
     {"A": "Vitamin A", "B": "Vitamin B12", "C": "Vitamin C", "D": "Vitamin D"}, "D"),

    ("DNA stands for:",
     {"A": "Deoxyribonucleic Acid", "B": "Diribonucleic Acid",
      "C": "Deoxyribose Nucleotide Acid", "D": "Double Nucleic Acid"}, "A"),

    ("Which organ filters blood in the human body?",
     {"A": "Liver", "B": "Lungs", "C": "Kidneys", "D": "Spleen"}, "C"),

    ("What is Newton's Second Law of Motion?",
     {"A": "F = mv", "B": "F = ma", "C": "E = mc²", "D": "P = mv"}, "B"),
]


# ─────────────────────────────────────────────
# Dosya oluşturucu
# ─────────────────────────────────────────────

def create_exam_docx(filename: str, title: str, questions: list) -> str:
    """
    Verilen soruları .docx formatında sınav dosyasına yazar.
    Döner: Kaydedilen dosyanın tam yolu.
    """
    doc = Document()

    # Başlık
    heading = doc.add_heading(title, level=1)
    heading.alignment = 1  # CENTER

    doc.add_paragraph(f"Total Questions: {len(questions)}")
    doc.add_paragraph("Choose the best answer for each question.")
    doc.add_paragraph()

    for i, (q_text, options, _) in enumerate(questions, 1):
        # Soru
        q_para = doc.add_paragraph()
        run = q_para.add_run(f"{i}. {q_text}")
        run.bold = True
        run.font.size = Pt(11)

        # Şıklar
        for letter in ["A", "B", "C", "D"]:
            if letter in options:
                opt_para = doc.add_paragraph()
                opt_para.paragraph_format.left_indent = Pt(20)
                opt_run = opt_para.add_run(f"{letter}) {options[letter]}")
                opt_run.font.size = Pt(10)

        doc.add_paragraph()  # Sorular arası boşluk

    filepath = str(OUTPUT_DIR / filename)
    doc.save(filepath)
    return filepath


def save_answer_key(filename: str, questions: list) -> str:
    """Cevap anahtarını JSON olarak kaydeder."""
    key = {
        str(i): {
            "answer": correct,
            "question": q_text[:60]
        }
        for i, (q_text, _, correct) in enumerate(questions, 1)
    }
    filepath = str(OUTPUT_DIR / filename)
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(key, f, indent=2, ensure_ascii=False)
    return filepath


if __name__ == "__main__":
    print("Örnek sınav dosyaları oluşturuluyor...\n")

    # Matematik
    math_path = create_exam_docx("math_exam.docx", "Mathematics MCQ Exam", MATH_QUESTIONS)
    math_key  = save_answer_key("math_exam_key.json", MATH_QUESTIONS)
    print(f"✓ Matematik sınavı  : {math_path}")
    print(f"  Cevap anahtarı    : {math_key}\n")

    # Fen
    sci_path = create_exam_docx("science_exam.docx", "Science MCQ Exam", SCIENCE_QUESTIONS)
    sci_key  = save_answer_key("science_exam_key.json", SCIENCE_QUESTIONS)
    print(f"✓ Fen sınavı        : {sci_path}")
    print(f"  Cevap anahtarı    : {sci_key}\n")

    print("Tamamlandı! Şimdi çalıştırabilirsin:")
    print("  python main.py sample_papers/math_exam.docx")
    print("  python main.py sample_papers/science_exam.docx")
