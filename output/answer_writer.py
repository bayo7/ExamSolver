"""
output/answer_writer.py
-----------------------
Çözülmüş cevapları düzenli bir .docx cevap kağıdına yazar.
Dosya adı kuralı: <orijinal_dosya_adı>_answers.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from typing import List, Dict, Any


# ─────────────────────────────────────────────
# Renk paleti
# ─────────────────────────────────────────────
COLOR_TITLE      = RGBColor(0x1A, 0x1A, 0x2E)  # Koyu lacivert
COLOR_CORRECT    = RGBColor(0x00, 0x7A, 0x3D)  # Koyu yeşil
COLOR_UNCERTAIN  = RGBColor(0xCC, 0x33, 0x00)  # Kırmızı
COLOR_HEADER_BG  = RGBColor(0x1A, 0x1A, 0x2E)  # Tablo başlık bg
COLOR_ALT_ROW    = RGBColor(0xF0, 0xF4, 0xFF)  # Alternatif satır
COLOR_BORDER     = RGBColor(0xCC, 0xCC, 0xCC)  # Tablo kenar


def set_cell_background(cell, hex_color: str):
    """Tablo hücresine arka plan rengi uygular."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    """Hücreye ince gri kenarlık uygular."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_heading(doc: Document, text: str, level: int = 1):
    """Belgeye stil uygulanmış başlık ekler."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_TITLE
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_TITLE
    return para


def write_answers(
    input_path: str,
    answers: List[Dict[str, Any]],
    output_dir: str = None
) -> str:
    """
    Ana çıktı fonksiyonu.
    Cevap listesinden düzenli bir .docx üretir.

    Parametreler:
        input_path  : Orijinal sınav dosyasının yolu
        answers     : solve_questions() çıktısı
        output_dir  : Çıktı klasörü (None → input ile aynı klasör)

    Döner:
        Oluşturulan dosyanın tam yolu
    """
    # Dosya adını belirle
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    out_name  = f"{base_name}_answers.docx"

    if output_dir is None:
        output_dir = os.path.dirname(os.path.abspath(input_path))
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, out_name)

    doc = Document()

    # ── Sayfa kenar boşlukları ──
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Başlık Bölümü ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("📋  ANSWER SHEET")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = COLOR_TITLE

    doc.add_paragraph()  # Boşluk

    # Meta bilgi tablosu
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.style = "Table Grid"

    meta_data = [
        ("Source File",   os.path.basename(input_path)),
        ("Generated",     datetime.now().strftime("%Y-%m-%d  %H:%M")),
    ]
    for i, (label, value) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            set_cell_border(cell)
            cell.paragraphs[0].runs[0].bold = (cell == row.cells[0])
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── İstatistik Özeti ──
    total     = len(answers)
    certain   = sum(1 for a in answers if a["answer"] != "UNCERTAIN")
    uncertain = total - certain
    avg_conf  = (
        sum(a["confidence"] for a in answers if a["answer"] != "UNCERTAIN") / certain
        if certain > 0 else 0
    )

    add_heading(doc, "Summary", level=2)

    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.style = "Table Grid"

    headers  = ["Total Q", "Answered", "Uncertain", "Avg Confidence"]
    values   = [str(total), str(certain), str(uncertain), f"{avg_conf:.0%}"]

    hdr_row = stats_table.rows[0]
    for i, (h, v) in enumerate(zip(headers, values)):
        cell = hdr_row.cells[i]
        cell.text = f"{h}\n{v}"
        set_cell_background(cell, "E8F4FD")
        set_cell_border(cell)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Cevap Tablosu ──
    add_heading(doc, "Answers", level=2)

    # Tablo başlıkları
    ans_table = doc.add_table(rows=1, cols=4)
    ans_table.style = "Table Grid"

    col_headers = ["#", "Question (Preview)", "Answer", "Confidence"]
    col_widths  = [Cm(1.2), Cm(9.0), Cm(2.5), Cm(2.5)]

    hdr_row = ans_table.rows[0]
    for i, (h, w) in enumerate(zip(col_headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        cell.text  = h
        set_cell_background(cell, "1A1A2E")
        set_cell_border(cell)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cevap satırları
    for idx, ans in enumerate(answers):
        row = ans_table.add_row()

        # Alternatif satır rengi
        bg_hex = "F0F4FF" if idx % 2 == 0 else "FFFFFF"

        # Sütun 1: Numara
        c0 = row.cells[0]
        c0.text = str(ans["question_number"])
        set_cell_background(c0, bg_hex)
        set_cell_border(c0)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c0.paragraphs[0].runs[0].font.size = Pt(9)

        # Sütun 2: Soru önizleme
        c1 = row.cells[1]
        preview = ans.get("question_text", "")[:90]
        if len(ans.get("question_text", "")) > 90:
            preview += "…"
        c1.text = preview
        set_cell_background(c1, bg_hex)
        set_cell_border(c1)
        c1.paragraphs[0].runs[0].font.size = Pt(9)

        # Sütun 3: Cevap
        c2 = row.cells[2]
        answer_val = ans["answer"]
        c2.text = answer_val
        set_cell_background(c2, bg_hex)
        set_cell_border(c2)
        run = c2.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_CORRECT if answer_val != "UNCERTAIN" else COLOR_UNCERTAIN
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sütun 4: Güven skoru
        c3 = row.cells[3]
        conf_str = f"{ans['confidence']:.0%}" if ans["answer"] != "UNCERTAIN" else "—"
        c3.text = conf_str
        set_cell_background(c3, bg_hex)
        set_cell_border(c3)
        c3.paragraphs[0].runs[0].font.size = Pt(9)
        c3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Detaylı Açıklamalar Bölümü ──
    add_heading(doc, "Detailed Reasoning", level=2)

    for ans in answers:
        # Soru başlığı
        q_para = doc.add_paragraph()
        q_run  = q_para.add_run(
            f"Q{ans['question_number']}.  →  {ans['answer']}"
        )
        q_run.bold = True
        q_run.font.size = Pt(10)
        q_run.font.color.rgb = COLOR_CORRECT if ans["answer"] != "UNCERTAIN" else COLOR_UNCERTAIN

        # Açıklama
        if ans.get("reason"):
            reason_para = doc.add_paragraph()
            reason_para.paragraph_format.left_indent = Cm(0.8)
            r_run = reason_para.add_run(ans["reason"])
            r_run.font.size = Pt(9)
            r_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # İnce ayraç çizgisi
        sep = doc.add_paragraph()
        sep_run = sep.add_run("─" * 60)
        sep_run.font.size = Pt(7)
        sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ── Footer notu ──
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "Generated by VisionSolve MCQ  •  AI-powered answer sheet"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    footer_run.italic = True

    doc.save(output_path)
    print(f"[Writer] Cevap kağıdı kaydedildi → {output_path}")
    return output_path


if __name__ == "__main__":
    # Test: Sahte cevaplarla çıktı üret
    dummy_answers = [
        {"question_number": 1, "question_text": "What is 2+2?",         "answer": "B", "confidence": 0.98, "reason": "Basic arithmetic."},
        {"question_number": 2, "question_text": "Capital of France?",   "answer": "A", "confidence": 0.99, "reason": "Paris is the capital."},
        {"question_number": 3, "question_text": "Unknown hard question", "answer": "UNCERTAIN", "confidence": 0.0, "reason": "API error"},
    ]
    path = write_answers("sample_exam.docx", dummy_answers, output_dir="output_answers")
    print(f"Test çıktısı: {path}")
